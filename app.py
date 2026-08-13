import streamlit as st
import sqlite3
from datetime import datetime, date
from zoneinfo import ZoneInfo
import pandas as pd
import math
import hashlib
import secrets
import io
import os
import re

# Optional image OCR. Install pytesseract + Pillow and the Tesseract OCR
# system package for automatic timetable-image reading.
try:
    from PIL import Image, ImageOps, ImageFilter
except Exception:
    Image = ImageOps = ImageFilter = None

try:
    import pytesseract
except Exception:
    pytesseract = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

st.set_page_config(
    page_title="Attendance Tracker",
    page_icon="📚",
    layout="centered",
)

DB_FILE = "attendance.db"
MIN_DATE = date(2026, 1, 1)
MAX_DATE = date(2099, 12, 31)
MAX_TIMETABLE_IMAGES = 7

DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
DAY_ALIASES = {
    "mon": "Monday", "monday": "Monday",
    "tue": "Tuesday", "tues": "Tuesday", "tuesday": "Tuesday",
    "wed": "Wednesday", "wednesday": "Wednesday",
    "thu": "Thursday", "thur": "Thursday", "thurs": "Thursday", "thursday": "Thursday",
    "fri": "Friday", "friday": "Friday",
    "sat": "Saturday", "saturday": "Saturday",
    "sun": "Sunday", "sunday": "Sunday",
}


def now_india():
    return datetime.now(ZoneInfo("Asia/Kolkata"))


def today_india():
    return now_india().date()


def hash_password(password, salt=None):
    salt = salt or secrets.token_bytes(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 200_000)
    return salt.hex() + ":" + digest.hex()


def verify_password(password, stored):
    try:
        salt_hex, digest_hex = stored.split(":")
        digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), bytes.fromhex(salt_hex), 200_000).hex()
        return secrets.compare_digest(digest, digest_hex)
    except Exception:
        return False


# =====================================================
# DATABASE + SAFE MIGRATIONS
# =====================================================

def get_connection():
    con = sqlite3.connect(DB_FILE, timeout=30)
    con.execute("PRAGMA journal_mode=WAL")

    con.execute("""CREATE TABLE IF NOT EXISTS students (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        full_name TEXT NOT NULL,
        roll_number TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        course TEXT DEFAULT '',
        division TEXT DEFAULT '',
        academic_year TEXT DEFAULT ''
    )""")

    con.execute("""CREATE TABLE IF NOT EXISTS attendance (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_name TEXT NOT NULL,
        roll_number TEXT NOT NULL,
        attendance_date TEXT NOT NULL,
        subject TEXT NOT NULL,
        status TEXT NOT NULL,
        UNIQUE(student_name, roll_number, attendance_date, subject)
    )""")

    con.execute("""CREATE TABLE IF NOT EXISTS timetable (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id INTEGER NOT NULL,
        day_name TEXT NOT NULL,
        lecture_no INTEGER NOT NULL,
        lecture_start TEXT DEFAULT '',
        lecture_end TEXT DEFAULT '',
        lecture_time TEXT DEFAULT '',
        subject_code TEXT DEFAULT '',
        subject_name TEXT NOT NULL,
        UNIQUE(student_id, day_name, lecture_no)
    )""")

    # Add columns to older versions without deleting existing data.
    cols = {r[1] for r in con.execute("PRAGMA table_info(students)").fetchall()}
    for col in ("course", "division", "academic_year"):
        if col not in cols:
            con.execute(f"ALTER TABLE students ADD COLUMN {col} TEXT DEFAULT ''")

    tcols = {r[1] for r in con.execute("PRAGMA table_info(timetable)").fetchall()}
    for col in ("lecture_start", "lecture_end"):
        if col not in tcols:
            con.execute(f"ALTER TABLE timetable ADD COLUMN {col} TEXT DEFAULT ''")

    # Date-specific daily subjects. If a row exists for a date, attendance uses it
    # instead of the normal weekly timetable for that date.
    con.execute("""CREATE TABLE IF NOT EXISTS daily_subjects (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id INTEGER NOT NULL,
        subject_date TEXT NOT NULL,
        lecture_no INTEGER NOT NULL,
        lecture_start TEXT DEFAULT '',
        lecture_end TEXT DEFAULT '',
        subject_code TEXT DEFAULT '',
        subject_name TEXT NOT NULL,
        UNIQUE(student_id, subject_date, lecture_no)
    )""")

    con.commit()
    return con


def register_student(name, roll, password):
    con = get_connection()
    try:
        con.execute("INSERT INTO students(full_name, roll_number, password_hash, course, division, academic_year) VALUES(?,?,?,?,?,?)",
                     (name, roll, hash_password(password), "", "", ""))
        con.commit()
        return True, "Account created successfully."
    except sqlite3.IntegrityError:
        return False, "This roll number is already registered."
    finally:
        con.close()


def login_student(roll, password):
    con = get_connection()
    row = con.execute("""SELECT id, full_name, roll_number, password_hash,
        COALESCE(course,''), COALESCE(division,''), COALESCE(academic_year,'')
        FROM students WHERE roll_number=?""", (roll,)).fetchone()
    con.close()
    if not row or not verify_password(password, row[3]):
        return None
    return {"id": row[0], "full_name": row[1], "roll_number": row[2],
            "course": row[4], "division": row[5], "academic_year": row[6]}


def update_profile(student_id, course, division, year):
    con = get_connection()
    con.execute("UPDATE students SET course=?, division=?, academic_year=? WHERE id=?",
                (course, division, year, student_id))
    con.commit(); con.close()


# =====================================================
# TIMETABLE
# =====================================================

def make_code(subject):
    code = re.sub(r"[^A-Za-z0-9]+", "", subject or "")[:12].upper()
    return code or "SUBJ"


def load_timetable(student_id):
    con = get_connection()
    rows = con.execute("""SELECT day_name, lecture_no,
        COALESCE(lecture_start,''), COALESCE(lecture_end,''),
        COALESCE(lecture_time,''), COALESCE(subject_code,''), subject_name
        FROM timetable WHERE student_id=? ORDER BY
        CASE day_name WHEN 'Monday' THEN 1 WHEN 'Tuesday' THEN 2 WHEN 'Wednesday' THEN 3
        WHEN 'Thursday' THEN 4 WHEN 'Friday' THEN 5 WHEN 'Saturday' THEN 6 WHEN 'Sunday' THEN 7 END,
        lecture_no""", (student_id,)).fetchall()
    con.close()
    result = {d: [] for d in DAYS}
    for r in rows:
        if r[0] in result:
            # (lecture_no, start, end, legacy_time, code, subject)
            result[r[0]].append(r[1:])
    return result


def save_timetable(student_id, rows):
    con = get_connection()
    try:
        con.execute("DELETE FROM timetable WHERE student_id=?", (student_id,))
        for r in rows:
            start = str(r.get("lecture_start", "")).strip()
            end = str(r.get("lecture_end", "")).strip()
            legacy = f"{start} - {end}" if start and end else str(r.get("lecture_time", "")).strip()
            con.execute("""INSERT INTO timetable
                (student_id, day_name, lecture_no, lecture_start, lecture_end, lecture_time, subject_code, subject_name)
                VALUES(?,?,?,?,?,?,?,?)""",
                (student_id, r["day_name"], int(r["lecture_no"]), start, end, legacy,
                 str(r.get("subject_code", "")).strip() or make_code(r["subject_name"]),
                 str(r["subject_name"]).strip()))
        con.commit(); return True, "Timetable saved successfully."
    except sqlite3.Error as e:
        con.rollback(); return False, f"Could not save timetable: {e}"
    finally:
        con.close()


def normalize_day(value):
    text = re.sub(r"\s+", " ", str(value).strip().lower())
    if text in DAY_ALIASES:
        return DAY_ALIASES[text]
    first = text.split()[0] if text.split() else ""
    return DAY_ALIASES.get(first, str(value).strip().title())


def clean_cell(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    text = str(value).strip()
    return "" if text.lower() in {"nan", "none", "nat"} else text


def find_column(columns, aliases):
    normalized = {str(c).strip().lower().replace("_", " "): c for c in columns}
    for alias in aliases:
        a = alias.lower().replace("_", " ")
        if a in normalized:
            return normalized[a]
    for col in columns:
        c = str(col).strip().lower().replace("_", " ")
        for alias in aliases:
            a = alias.lower().replace("_", " ")
            if a in c or c in a:
                return col
    return None


def split_time_range(value):
    text = clean_cell(value)
    if not text:
        return "", ""
    # Handles 09:00-10:00, 9.00 to 10.00, 9 AM - 10 AM, etc.
    pat = re.compile(r"(\d{1,2}(?:[:.]\d{2})?\s*(?:am|pm)?)\s*(?:-|–|—|to)\s*(\d{1,2}(?:[:.]\d{2})?\s*(?:am|pm)?)", re.I)
    m = pat.search(text)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return text, ""


def dataframe_to_timetable(df):
    if df is None or df.empty:
        return None, "No timetable data was found."
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    day_col = find_column(df.columns, ["day", "day name", "weekday", "week day"])
    subject_col = find_column(df.columns, ["subject", "subject name", "class", "course", "paper", "lecture subject"])
    lecture_col = find_column(df.columns, ["lecture", "lecture no", "lecture number", "period", "period no", "period number", "slot"])
    start_col = find_column(df.columns, ["lecture start", "start time", "start", "from"])
    end_col = find_column(df.columns, ["lecture end", "end time", "end", "to"])
    time_col = find_column(df.columns, ["time", "lecture time", "timing", "time slot", "period time"])
    code_col = find_column(df.columns, ["subject code", "code", "subject_code", "paper code"])

    if not day_col or not subject_col:
        return None, "I could not find Day and Subject columns. You can use Daily Subjects or upload a clearer table."

    rows = []
    counters = {d: 0 for d in DAYS}
    for _, row in df.iterrows():
        day = normalize_day(row.get(day_col, ""))
        subject = clean_cell(row.get(subject_col, ""))
        if day not in DAYS or not subject:
            continue
        raw_no = clean_cell(row.get(lecture_col, "")) if lecture_col else ""
        try:
            no = int(float(raw_no)) if raw_no else None
        except Exception:
            m = re.search(r"\d+", raw_no); no = int(m.group()) if m else None
        if not no:
            counters[day] += 1; no = counters[day]
        else:
            counters[day] = max(counters[day], no)
        start = clean_cell(row.get(start_col, "")) if start_col else ""
        end = clean_cell(row.get(end_col, "")) if end_col else ""
        if not start and time_col:
            start, end = split_time_range(row.get(time_col, ""))
        code = clean_cell(row.get(code_col, "")) if code_col else ""
        rows.append({"day_name": day, "lecture_no": no, "lecture_start": start,
                     "lecture_end": end, "subject_code": code or make_code(subject), "subject_name": subject})
    if not rows:
        return None, "No valid timetable rows were found."
    used = {d: set() for d in DAYS}; final = []
    for r in rows:
        no = r["lecture_no"]
        while no in used[r["day_name"]]: no += 1
        r["lecture_no"] = no; used[r["day_name"]].add(no); final.append(r)
    final.sort(key=lambda r: (DAYS.index(r["day_name"]), r["lecture_no"]))
    return final, None


def parse_text_timetable(text):
    rows = []; counters = {d: 0 for d in DAYS}
    for line in [x.strip() for x in text.splitlines() if x.strip()]:
        parts = [p.strip() for p in re.split(r"\t+|\s{2,}|\|+|;", line) if p.strip()]
        if len(parts) < 2: continue
        day = normalize_day(parts[0])
        if day not in DAYS: continue
        remaining = parts[1:]
        no = None; start = end = ""; code = ""; subject = ""
        for part in remaining:
            if no is None:
                m = re.fullmatch(r"(?:lecture|period|slot)?\s*(\d+)", part, re.I)
                if m: no = int(m.group(1)); continue
            if not start:
                s, e = split_time_range(part)
                if s: start, end = s, e; continue
            if not code and re.fullmatch(r"[A-Za-z]{1,12}[-_/]?\d{0,5}", part):
                code = part; continue
            if not subject: subject = part
        if not subject and remaining: subject = remaining[-1]
        if not subject: continue
        if no is None: counters[day] += 1; no = counters[day]
        rows.append({"day_name": day, "lecture_no": no, "lecture_start": start, "lecture_end": end,
                     "subject_code": code or make_code(subject), "subject_name": subject})
    return (rows, None) if rows else (None, "No timetable rows could be detected from this text.")


def ocr_image(image_bytes):
    if Image is None or pytesseract is None:
        return "", "Image OCR is not installed. Install Pillow and pytesseract plus the Tesseract OCR package."
    try:
        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        gray = ImageOps.grayscale(image)
        gray = ImageOps.autocontrast(gray)
        gray = gray.resize((max(1, gray.width * 2), max(1, gray.height * 2)))
        gray = gray.filter(ImageFilter.SHARPEN)
        text = pytesseract.image_to_string(gray, config="--psm 6")
        return text, None
    except Exception as e:
        return "", f"Could not OCR image: {e}"


def parse_image_timetable(image_bytes, forced_day=""):
    text, error = ocr_image(image_bytes)
    if error:
        return [], error, text
    detected_day = forced_day
    if not detected_day:
        low = text.lower()
        for d in DAYS:
            if re.search(r"\b" + re.escape(d.lower()) + r"\b", low):
                detected_day = d; break
    if not detected_day:
        return [], "Day was not detected. Select the day for this image and try again.", text
    # Convert OCR lines into a tiny table for the common one-row-per-lecture case.
    rows = []
    counters = 0
    for line in [x.strip() for x in text.splitlines() if x.strip()]:
        # Skip obvious title/header lines.
        if line.lower() in {d.lower() for d in DAYS}: continue
        times = re.search(r"(\d{1,2}(?:[:.]\d{2})?\s*(?:am|pm)?)\s*(?:-|–|—|to)\s*(\d{1,2}(?:[:.]\d{2})?\s*(?:am|pm)?)", line, re.I)
        if times:
            start, end = times.group(1).strip(), times.group(2).strip()
            subject = line[:times.start()] + " " + line[times.end():]
            subject = re.sub(r"^[\s\-:|\d.]+|[\s\-:|]+$", "", subject).strip()
            subject = re.sub(r"\s+", " ", subject)
            if subject and len(subject) > 1:
                counters += 1
                rows.append({"day_name": detected_day, "lecture_no": counters,
                             "lecture_start": start, "lecture_end": end,
                             "subject_code": make_code(subject), "subject_name": subject})
    if not rows:
        # Fallback: use non-empty OCR lines as subjects; user can correct them.
        candidates = []
        for line in text.splitlines():
            line = re.sub(r"\s+", " ", line).strip()
            if len(line) >= 2 and not re.search(r"\b(?:monday|tuesday|wednesday|thursday|friday|saturday|sunday)\b", line, re.I):
                if not re.fullmatch(r"[\d\W_]+", line): candidates.append(line)
        for subject in candidates[:30]:
            counters += 1
            rows.append({"day_name": detected_day, "lecture_no": counters,
                         "lecture_start": "", "lecture_end": "",
                         "subject_code": make_code(subject), "subject_name": subject})
    return rows, None, text


def read_timetable_file(upload):
    filename = (upload.name or "").lower(); ext = os.path.splitext(filename)[1]
    try:
        raw = upload.getvalue()
        if ext in {".csv", ".tsv"}:
            df = pd.read_csv(io.BytesIO(raw), sep="\t" if ext == ".tsv" else ",")
            return dataframe_to_timetable(df)
        if ext in {".xlsx", ".xls", ".xlsm"}:
            sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
            if not sheets: return None, "No worksheet was found."
            return dataframe_to_timetable(max(sheets.values(), key=lambda x: len(x)))
        if ext == ".json":
            data = json.loads(raw.decode("utf-8-sig"))
            if isinstance(data, list): return dataframe_to_timetable(pd.DataFrame(data))
            if isinstance(data, dict):
                items=[]
                for day, vals in data.items():
                    d=normalize_day(day)
                    if d not in DAYS or not isinstance(vals,list): continue
                    for i,v in enumerate(vals,1):
                        if isinstance(v,dict):
                            items.append({"Day":d,"Lecture":v.get("Lecture",i),"Lecture Start":v.get("Lecture Start",v.get("start","")),"Lecture End":v.get("Lecture End",v.get("end","")),"Subject":v.get("Subject",v.get("subject","")),"Subject Code":v.get("Subject Code",v.get("code",""))})
                        else: items.append({"Day":d,"Lecture":i,"Subject":str(v)})
                return dataframe_to_timetable(pd.DataFrame(items)) if items else (None,"JSON timetable format not recognized.")
        if ext in {".txt", ".text"}: return parse_text_timetable(raw.decode("utf-8", errors="ignore"))
        if ext in {".html", ".htm"}: return dataframe_to_timetable(max(pd.read_html(io.BytesIO(raw)), key=lambda x: len(x)))
        if ext == ".pdf":
            if PdfReader is None: return None, "PDF reader is not installed."
            text="\n".join(p.extract_text() or "" for p in PdfReader(io.BytesIO(raw)).pages)
            return parse_text_timetable(text) if text.strip() else (None,"Scanned/image PDFs need image upload or manual entry.")
        if ext == ".docx":
            if Document is None: return None,"DOCX reader is not installed."
            doc=Document(io.BytesIO(raw)); frames=[]
            for table in doc.tables:
                data=[[c.text.strip() for c in row.cells] for row in table.rows]
                if len(data)>=2: frames.append(pd.DataFrame(data[1:],columns=data[0]))
            if frames: return dataframe_to_timetable(max(frames,key=lambda x:len(x)))
            return parse_text_timetable("\n".join(p.text for p in doc.paragraphs if p.text.strip()))
        return None, "This file can be uploaded, but automatic timetable reading is available for CSV/Excel/TXT/JSON/HTML/PDF/DOCX. For timetable photos use the Image Upload section."
    except Exception as e:
        return None, f"Could not read timetable file: {e}"


# =====================================================
# DAILY SUBJECTS
# =====================================================

def load_daily_subjects(student_id, selected_date):
    con=get_connection()
    rows=con.execute("""SELECT lecture_no, COALESCE(lecture_start,''), COALESCE(lecture_end,''),
        COALESCE(subject_code,''), subject_name FROM daily_subjects
        WHERE student_id=? AND subject_date=? ORDER BY lecture_no""",
        (student_id, selected_date.isoformat())).fetchall()
    con.close(); return rows


def save_daily_subjects(student_id, selected_date, rows):
    con=get_connection()
    try:
        con.execute("DELETE FROM daily_subjects WHERE student_id=? AND subject_date=?", (student_id, selected_date.isoformat()))
        for i,r in enumerate(rows,1):
            subject=str(r.get("subject_name","")).strip()
            if not subject: continue
            con.execute("""INSERT INTO daily_subjects(student_id,subject_date,lecture_no,lecture_start,lecture_end,subject_code,subject_name)
                VALUES(?,?,?,?,?,?,?)""", (student_id,selected_date.isoformat(),i,str(r.get("lecture_start","")).strip(),str(r.get("lecture_end","")).strip(),str(r.get("subject_code","")).strip() or make_code(subject),subject))
        con.commit(); return True,"Daily subjects saved."
    except sqlite3.Error as e:
        con.rollback(); return False,f"Could not save daily subjects: {e}"
    finally: con.close()


def classes_for_date(student_id, selected_date):
    daily=load_daily_subjects(student_id,selected_date)
    if daily:
        return [(r[0],r[1],r[2],r[3],r[4]) for r in daily]
    weekly=load_timetable(student_id).get(selected_date.strftime("%A"),[])
    return [(r[0],r[1],r[2],r[4],r[5]) for r in weekly]


# =====================================================
# ATTENDANCE
# =====================================================

def save_attendance(name, roll, selected_date, attendance_data):
    con=get_connection()
    try:
        for subject,status in attendance_data.items():
            con.execute("""INSERT INTO attendance(student_name,roll_number,attendance_date,subject,status)
                VALUES(?,?,?,?,?) ON CONFLICT(student_name,roll_number,attendance_date,subject)
                DO UPDATE SET status=excluded.status""", (name,roll,selected_date.isoformat(),subject,status))
        con.commit(); return True,"Attendance saved successfully."
    except sqlite3.Error as e:
        con.rollback(); return False,f"Could not save attendance: {e}"
    finally: con.close()


def load_attendance(name, roll):
    con=get_connection(); df=pd.read_sql_query("SELECT attendance_date,subject,status FROM attendance WHERE student_name=? AND roll_number=? ORDER BY attendance_date DESC",con,params=(name,roll)); con.close(); return df


def needed_for_75(present,total):
    if total<=0 or present/total>=.75: return 0
    return max(0,math.ceil((.75*total-present)/.25))


def can_miss_at_75(present,total):
    if total<=0 or present/total<.75: return 0
    return max(0,math.floor(present/.75-total))


# =====================================================
# START APP
# =====================================================
get_connection().close()
DEFAULTS={"logged_in":False,"student_id":None,"student_name":"","roll_number":"","course":"","division":"","academic_year":""}
for k,v in DEFAULTS.items():
    if k not in st.session_state: st.session_state[k]=v

if not st.session_state.logged_in:
    st.title("📚 Attendance Tracker")
    st.caption("For any student • any course • any division • any academic year")
    login_tab,create_tab=st.tabs(["🔐 Login","📝 Create Account"])
    with login_tab:
        st.subheader("Login to your account")
        roll=st.text_input("Roll Number",key="login_roll")
        password=st.text_input("Password",type="password",key="login_password")
        if st.button("🔐 Login",type="primary",width="stretch"):
            student=login_student(roll.strip(),password)
            if student:
                st.session_state.logged_in=True; st.session_state.student_id=student["id"]; st.session_state.student_name=student["full_name"]; st.session_state.roll_number=student["roll_number"]; st.session_state.course=student["course"]; st.session_state.division=student["division"]; st.session_state.academic_year=student["academic_year"]; st.rerun()
            else: st.error("Invalid roll number or password.")
    with create_tab:
        st.subheader("Create your account")
        name=st.text_input("Full Name",key="create_name")
        roll_new=st.text_input("Roll Number",key="create_roll")
        p1=st.text_input("Create Password",type="password",key="create_password")
        p2=st.text_input("Confirm Password",type="password",key="confirm_password")
        if st.button("📝 Create Account",type="primary",width="stretch"):
            name=name.strip(); roll_new=roll_new.strip()
            if not name or not roll_new or not p1: st.error("Please fill all required fields.")
            elif len(p1)<6: st.error("Password must be at least 6 characters.")
            elif p1!=p2: st.error("Passwords do not match.")
            else:
                ok,msg=register_student(name,roll_new,p1); (st.success if ok else st.error)(msg)
    st.stop()

student_id=st.session_state.student_id; student_name=st.session_state.student_name; roll_number=st.session_state.roll_number
st.title("📚 Attendance Tracker")
parts=[x for x in [st.session_state.course,st.session_state.division,st.session_state.academic_year] if x]
st.caption(" • ".join(parts) if parts else "Set up your class details below.")

with st.sidebar:
    st.header("👤 My Account")
    st.write(f"**Name:** {student_name}"); st.write(f"**Roll No.:** {roll_number}")
    if st.session_state.course: st.write(f"**Class:** {st.session_state.course}")
    if st.session_state.division: st.write(f"**Division:** {st.session_state.division}")
    if st.session_state.academic_year: st.write(f"**Academic Year:** {st.session_state.academic_year}")
    if st.button("🚪 Logout",width="stretch"):
        for k,v in DEFAULTS.items(): st.session_state[k]=v
        st.rerun()

setup_tab,daily_tab,attendance_tab,dashboard_tab,history_tab=st.tabs(["⚙️ Setup","📅 Daily Subjects","📝 Attendance","📊 Dashboard","📚 History"])

# SETUP
with setup_tab:
    st.header("⚙️ Class & Timetable Setup")
    st.info("Nothing is hard-coded to B.Com, M-2, FY B.Com, 2026-27, IKS-HDA, or any particular weekday. Every student chooses their own class details.")
    with st.form("profile"):
        course=st.text_input("Course / Class",value=st.session_state.course,placeholder="B.Com / BCA / BBA / Class 12 / Any Class")
        division=st.text_input("Division / Section",value=st.session_state.division,placeholder="M-2 / A / B / Any Section")
        year=st.text_input("Academic Year",value=st.session_state.academic_year,placeholder="2026-27 / 2027-28 / 2028-29 / Any")
        if st.form_submit_button("💾 Save Class Details",type="primary",width="stretch"):
            update_profile(student_id,course.strip(),division.strip(),year.strip()); st.session_state.course=course.strip(); st.session_state.division=division.strip(); st.session_state.academic_year=year.strip(); st.success("Class details saved."); st.rerun()

    st.divider(); st.subheader("🖼️ Upload Timetable Images")
    st.write(f"Upload up to **{MAX_TIMETABLE_IMAGES} images** of your timetable. JPG, JPEG, PNG, WEBP, BMP and GIF are supported.")
    st.caption("The app will OCR the image, detect days/times/subjects where possible, and show an editable preview before saving. OCR is best-effort, so always check the preview.")
    images=st.file_uploader("Choose timetable images",type=["jpg","jpeg","png","webp","bmp","gif"],accept_multiple_files=True,key="timetable_images")
    if images and len(images)>MAX_TIMETABLE_IMAGES:
        st.error(f"Maximum {MAX_TIMETABLE_IMAGES} timetable images allowed. Please remove {len(images)-MAX_TIMETABLE_IMAGES} image(s).")
        images=images[:MAX_TIMETABLE_IMAGES]
    image_rows=[]
    if images:
        for idx,up in enumerate(images,1):
            st.markdown(f"**Image {idx}: {up.name}**")
            day_choice=st.selectbox("Day for this image (choose only if the day is not detected automatically)", ["Auto-detect"]+DAYS, key=f"img_day_{idx}")
            with st.expander("Preview image",expanded=False): st.image(up, use_container_width=True)
            forced="" if day_choice=="Auto-detect" else day_choice
            rows,error,ocr_text=parse_image_timetable(up.getvalue(),forced)
            if error: st.warning(error)
            if ocr_text:
                with st.expander("OCR text (for checking)",expanded=False): st.text(ocr_text[:12000])
            image_rows.extend(rows)
        if image_rows:
            st.subheader("🔎 Detected timetable — check and edit before saving")
            idf=pd.DataFrame(image_rows)
            idf=idf[["day_name","lecture_no","lecture_start","lecture_end","subject_code","subject_name"]]
            edited=st.data_editor(idf,use_container_width=True,num_rows="dynamic",hide_index=True,
                                  column_config={"day_name":st.column_config.SelectboxColumn("Day",options=DAYS),"lecture_no":st.column_config.NumberColumn("Lecture",min_value=1,step=1),"lecture_start":"Lecture Start","lecture_end":"Lecture End","subject_code":"Subject Code","subject_name":"Subject"},key="ocr_editor")
            if st.button("💾 Save Detected Timetable",type="primary",width="stretch"):
                rows=[]
                for _,r in edited.iterrows():
                    subject=clean_cell(r.get("subject_name")); day=normalize_day(r.get("day_name"))
                    if subject and day in DAYS:
                        rows.append({"day_name":day,"lecture_no":int(r.get("lecture_no") or 1),"lecture_start":clean_cell(r.get("lecture_start")),"lecture_end":clean_cell(r.get("lecture_end")),"subject_code":clean_cell(r.get("subject_code")) or make_code(subject),"subject_name":subject})
                ok,msg=save_timetable(student_id,rows); (st.success if ok else st.error)(msg)
                if ok: st.rerun()

    st.divider(); st.subheader("📄 Optional CSV / Excel timetable")
    st.caption("Image upload is the main method. CSV/Excel can also be used if someone already has a table; PDF upload is not required.")
    upload=st.file_uploader("Upload CSV or Excel timetable",type=["csv","tsv","xlsx","xls","xlsm"],key="table_upload")
    if upload:
        rows,error=read_timetable_file(upload)
        if error: st.warning(error)
        if rows:
            preview=pd.DataFrame(rows)
            st.dataframe(preview.rename(columns={"day_name":"Day","lecture_no":"Lecture","lecture_start":"Lecture Start","lecture_end":"Lecture End","subject_code":"Subject Code","subject_name":"Subject"}),use_container_width=True,hide_index=True)
            if st.button("💾 Save Table Timetable",type="primary",width="stretch"):
                ok,msg=save_timetable(student_id,rows); (st.success if ok else st.error)(msg)
                if ok: st.rerun()

    st.divider(); st.subheader("✏️ Weekly Timetable Editor")
    st.caption("Use this to correct OCR results or create a timetable manually. There is no limit of 3 or 4 lectures: add as many as your class needs.")
    current=load_timetable(student_id); manual=[]
    for day in DAYS:
        with st.expander(day,expanded=(day=="Monday")):
            old=current.get(day,[])
            count=st.number_input(f"Lectures on {day}",min_value=0,max_value=50,value=len(old),step=1,key=f"wk_count_{day}")
            for i in range(int(count)):
                o=old[i] if i<len(old) else (i+1,"","","","","")
                c1,c2,c3,c4=st.columns([1,1,1.2,2.4])
                with c1: start=st.text_input("Lecture Start",value=str(o[1]),key=f"wk_start_{day}_{i}")
                with c2: end=st.text_input("Lecture End",value=str(o[2]),key=f"wk_end_{day}_{i}")
                with c3: code=st.text_input("Subject Code",value=str(o[4]),key=f"wk_code_{day}_{i}")
                with c4: subject=st.text_input("Subject",value=str(o[5]),key=f"wk_sub_{day}_{i}")
                if subject.strip(): manual.append({"day_name":day,"lecture_no":i+1,"lecture_start":start.strip(),"lecture_end":end.strip(),"subject_code":code.strip() or make_code(subject),"subject_name":subject.strip()})
    if st.button("💾 Save Weekly Timetable",type="primary",width="stretch"):
        ok,msg=save_timetable(student_id,manual); (st.success if ok else st.error)(msg)
        if ok: st.rerun()

# DAILY SUBJECTS
with daily_tab:
    st.header("📅 Daily Subjects")
    st.write("Choose any date. The app automatically loads that day's subjects from the weekly timetable. You can then add, remove, or change lectures for that specific date.")
    selected_daily=st.date_input("Date",value=today_india(),min_value=MIN_DATE,max_value=MAX_DATE,key="daily_date")
    day_name=selected_daily.strftime("%A")
    existing=load_daily_subjects(student_id,selected_daily)
    base=existing if existing else load_timetable(student_id).get(day_name,[])
    if existing: st.success("This date has a custom daily subject list.")
    else: st.info(f"Using the normal {day_name} timetable. Saving here creates a date-specific list.")
    n=st.number_input("Number of lectures for this date",min_value=0,max_value=50,value=len(base),step=1,key=f"daily_count_{selected_daily}")
    daily_rows=[]
    for i in range(int(n)):
        if i<len(base):
            if len(base[i])==5: o=(base[i][0],base[i][1],base[i][2],base[i][3],base[i][4])
            else: o=(base[i][0],base[i][1],base[i][2],base[i][4],base[i][5])
        else: o=(i+1,"","","","")
        c1,c2,c3,c4=st.columns([1,1,1.2,2.4])
        with c1: s=st.text_input("Lecture Start",value=str(o[1]),key=f"d_start_{selected_daily}_{i}")
        with c2: e=st.text_input("Lecture End",value=str(o[2]),key=f"d_end_{selected_daily}_{i}")
        with c3: code=st.text_input("Subject Code",value=str(o[3]),key=f"d_code_{selected_daily}_{i}")
        with c4: subj=st.text_input("Subject",value=str(o[4]),key=f"d_sub_{selected_daily}_{i}")
        if subj.strip(): daily_rows.append({"subject_name":subj.strip(),"lecture_start":s.strip(),"lecture_end":e.strip(),"subject_code":code.strip() or make_code(subj)})
    if st.button("💾 Save Daily Subjects",type="primary",width="stretch"):
        ok,msg=save_daily_subjects(student_id,selected_daily,daily_rows); (st.success if ok else st.error)(msg)
        if ok: st.rerun()
    st.caption("If you want to go back to the normal weekly timetable for this date, save 0 lectures for the date.")

# ATTENDANCE
with attendance_tab:
    st.header("📝 Mark Attendance")
    selected=st.date_input("Select Date",value=today_india(),min_value=MIN_DATE,max_value=MAX_DATE,key="attendance_date")
    day=selected.strftime("%A")
    st.subheader(f"📅 {day}, {selected.strftime('%d %B %Y')}")
    if selected>today_india():
        st.warning("⚠️ This is a future date. Attendance can be recorded, but check the date carefully.")
    classes=classes_for_date(student_id,selected)
    if not classes:
        st.warning("No subjects are configured for this date. Add them in Daily Subjects or Setup.")
    else:
        holiday=st.checkbox("🟡 College Holiday — whole day",key=f"holiday_{selected}")
        attendance={}
        for no,start,end,code,subj in classes:
            st.divider(); st.markdown(f"### Lecture {no} — {start} → {end}" if start or end else f"### Lecture {no}"); st.write(f"**{code} — {subj}**")
            key=f"{code} — {subj}"
            attendance[key]="Holiday" if holiday else st.radio("Status",["Present","Absent","Holiday"],horizontal=True,key=f"status_{selected}_{no}_{code}_{subj}")
        if st.button("💾 Save Attendance",type="primary",width="stretch"):
            ok,msg=save_attendance(student_name,roll_number,selected,attendance); (st.success if ok else st.error)(msg)
            if ok: st.rerun()

# DASHBOARD
with dashboard_tab:
    st.header("📊 Attendance Dashboard")
    data=load_attendance(student_name,roll_number)
    if data.empty: st.info("No attendance recorded yet.")
    else:
        conducted=data[data.status!="Holiday"]; present=len(data[data.status=="Present"]); absent=len(data[data.status=="Absent"]); total=len(conducted); overall=present/total*100 if total else 0
        a,b,c=st.columns(3); a.metric("Overall",f"{overall:.1f}%"); b.metric("Present",present); c.metric("Absent",absent)
        subjects=[]; tt=load_timetable(student_id)
        for d in DAYS:
            for r in tt.get(d,[]):
                label=f"{r[4]} — {r[5]}"
                if label not in subjects: subjects.append(label)
        for v in data.subject.dropna().unique():
            if v not in subjects: subjects.append(v)
        report=[]
        for subject in subjects:
            d=data[data.subject==subject]; p=len(d[d.status=="Present"]); ab=len(d[d.status=="Absent"]); h=len(d[d.status=="Holiday"]); t=p+ab; pct=p/t*100 if t else 0
            report.append({"Subject":subject,"Present":p,"Absent":ab,"Holiday":h,"Total":t,"Attendance %":round(pct,1),"Need for 75%":needed_for_75(p,t),"Can Miss":can_miss_at_75(p,t)})
        st.subheader("📚 Subject-wise Attendance"); st.dataframe(pd.DataFrame(report),use_container_width=True,hide_index=True)
        st.subheader("🎯 75% Attendance Planner"); planner=[]
        for r in report:
            planner.append({"Subject":r["Subject"],"Current %":r["Attendance %"],"Present":r["Present"],"Total Conducted":r["Total"],"Need to Reach 75%":r["Need for 75%"],"Can Miss at 75%":r["Can Miss"]})
        st.dataframe(pd.DataFrame(planner),use_container_width=True,hide_index=True)

# HISTORY
with history_tab:
    st.header("📚 Attendance History")
    data=load_attendance(student_name,roll_number)
    if data.empty: st.info("No attendance history yet.")
    else:
        history=data.copy(); history["attendance_date"]=pd.to_datetime(history["attendance_date"]).dt.strftime("%d-%m-%Y"); history.columns=["Date","Subject","Status"]
        st.dataframe(history,use_container_width=True,hide_index=True)
        st.download_button("⬇️ Download Attendance CSV",history.to_csv(index=False).encode("utf-8"),"attendance.csv","text/csv",width="stretch")
